"""
create_template.py
------------------
One-time script to generate templates/valuation_report.docx.
Run this once with:
    python create_template.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top, bottom, left, right  → color hex."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, color in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_styled_heading(doc, text: str, level: int, color_hex="1B4F72"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return p


def add_kv_table(doc, rows: list[tuple[str, str]]):
    """Add a 2-column key/value table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        row = table.rows[i]
        # Key cell
        kc = row.cells[0]
        kc.width = Inches(2.2)
        set_cell_bg(kc, "D6EAF8")
        kp = kc.paragraphs[0]
        kp.clear()
        kr = kp.add_run(key)
        kr.bold = True
        kr.font.size = Pt(10)
        # Value cell
        vc = row.cells[1]
        vp = vc.paragraphs[0]
        vp.clear()
        vp.add_run(val)
        vp.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def build_template():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Header bar ────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("REAL ESTATE VALUATION REPORT")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Prepared for: {{ address }}, {{ city }}, {{ state }} {{ zip_code }}")
    sr.font.size = Pt(11)
    sr.italic = True

    doc.add_paragraph()

    # ── Report metadata table ─────────────────────────────────────────────────
    add_kv_table(doc, [
        ("Report Date",    "{{ report_date }}"),
        ("Property Address", "{{ address }}, {{ city }}, {{ state }} {{ zip_code }}"),
        ("List Price",     "{{ list_price_display }}"),
        ("Prepared By",    "Antigravity Valuation Engine"),
    ])

    # ── Section 1: Subject Property Details ──────────────────────────────────
    add_styled_heading(doc, "1. Subject Property Details", level=2)
    add_kv_table(doc, [
        ("Square Footage",  "{{ sq_ft }} sq ft"),
        ("Bedrooms",        "{{ bedrooms }}"),
        ("Bathrooms",       "{{ bathrooms }}"),
        ("Lot Size",        "{{ lot_size }}"),
        ("Year Built",      "{{ year_built }}"),
        ("Price / Sq Ft",   "{{ price_per_sqft }}"),
    ])

    # ── Section 2: Comparable Sales ───────────────────────────────────────────
    add_styled_heading(doc, "2. Comparable Sales", level=2)
    comp_table = doc.add_table(rows=4, cols=4)
    comp_table.style = "Table Grid"
    headers = ["Comp #", "Address", "Sale Price", "Sq Ft", ]
    hdr_row = comp_table.rows[0]
    # Add extra col for $/sqft
    comp_table.add_column(Inches(1.0))
    headers.append("$/Sq Ft")
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        set_cell_bg(cell, "1B4F72")
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    for i in range(1, 4):
        row = comp_table.rows[i]
        vals = [
            str(i),
            f"{{{{ comp{i}_address }}}}",
            f"{{{{ comp{i}_price_display }}}}",
            f"{{{{ comp{i}_sqft_display }}}}",
            f"{{{{ comp{i}_ppsf }}}}",
        ]
        for j, v in enumerate(vals):
            cell = row.cells[j]
            if i % 2 == 0:
                set_cell_bg(cell, "EAF4FB")
            p = cell.paragraphs[0]
            p.clear()
            p.add_run(v).font.size = Pt(10)

    doc.add_paragraph()

    # ── Section 3: Valuation Summary ──────────────────────────────────────────
    add_styled_heading(doc, "3. Valuation Summary", level=2)
    add_kv_table(doc, [
        ("Avg Comp $/Sq Ft",   "{{ avg_comp_ppsf }}"),
        ("Subject $/Sq Ft",    "{{ price_per_sqft }}"),
        ("Estimated Market Value", "{{ estimated_value }}"),
    ])

    # ── Notes / Disclaimer ────────────────────────────────────────────────────
    add_styled_heading(doc, "Notes & Disclaimer", level=3)
    disclaimer = doc.add_paragraph(
        "This report is generated automatically by the Antigravity Valuation Engine "
        "for informational purposes only. Values are estimates based on comparable "
        "sales data and should not be used as a substitute for a certified appraisal."
    )
    disclaimer.runs[0].font.size = Pt(9)
    disclaimer.runs[0].italic = True

    # ── Footer rule ───────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph("─" * 80)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    os.makedirs("templates", exist_ok=True)
    doc.save("templates/valuation_report.docx")
    print("✅ Template saved to templates/valuation_report.docx")


if __name__ == "__main__":
    build_template()
